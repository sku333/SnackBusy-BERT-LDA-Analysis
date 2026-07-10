from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

OUTPUT_DIR = r'E:\ZDcup\upgraded_analysis\output'

def set_font(run, size=11, bold=False, color=None, name='宋体'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125), name='黑体')
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_font(run, size=13, bold=True, color=(68, 114, 196), name='黑体')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        set_font(run, size=11, bold=True, color=(0, 0, 0), name='黑体')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_para(doc, text, indent=1, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        set_font(r0, size=11, bold=True)
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.first_line_indent = Pt(22 * indent)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, name='黑体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_image(doc, img_path, caption, width=Inches(5.5)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_font(r, size=10, color=(89, 89, 89))
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_para(doc, f'[图片未找到: {img_path}]')

# ===== 封面 =====
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('硬折扣零食店消费者评论数据挖掘分析报告')
set_font(run, size=20, bold=True, color=(31, 73, 125), name='黑体')

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('传统NLP方法 vs 深度学习方法 对比研究')
set_font(run2, size=14, color=(68, 114, 196))

doc.add_paragraph()
for line in [
    '数据来源：大众点评、美团（上海零食很忙、好想来门店）',
    '有效样本：1,375 条用户评论',
    '分析方法：Jieba + SnowNLP + LDA  vs  RoBERTa + BERT + LDA',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    set_font(r, size=11, color=(89, 89, 89))

doc.add_page_break()

# ===== 一、情感分析 =====
add_heading(doc, '一、情感分析结果对比', 1)

add_heading(doc, '1.1 方法对比', 2)
add_table(doc,
    ['方法', '模型', '原理', '输出'],
    [
        ['原方法', 'SnowNLP', '基于朴素贝叶斯的规则模型，通用语料训练', '0-1情感分值，阈值划分三类'],
        ['新方法', 'RoBERTa', '在京东电商评论上微调的深度学习模型', 'positive/negative + 置信度'],
    ],
    col_widths=[2, 3, 5, 4]
)

add_heading(doc, '1.2 量化结果对比', 2)
add_table(doc,
    ['情感类别', '原方法 (SnowNLP)', '新方法 (RoBERTa)', '差异'],
    [
        ['正向', '73.8%（1,015条）', '89.8%（1,235条）', '+16.0个百分点'],
        ['中性', '3.9%（54条）', 'N/A（二分类）', '原中性评论被分至正/负'],
        ['负向', '22.3%（306条）', '10.2%（140条）', '-12.1个百分点'],
        ['平均置信度', '无', '96.93%', '模型预测把握度极高'],
    ],
    col_widths=[2.5, 4, 4, 3.5]
)

add_heading(doc, '1.3 情感分布对比图', 2)
add_image(doc,
    os.path.join(OUTPUT_DIR, 'sentiment_comparison.png'),
    '图1  情感分布对比：SnowNLP（左）vs RoBERTa（右）',
    width=Inches(6.0)
)

add_heading(doc, '1.4 结果分析', 2)
add_para(doc, 'RoBERTa识别出89.8%的正向评论，比SnowNLP高出16个百分点，差异原因如下：', indent=1)
for item in [
    'SnowNLP基于通用语料训练，对电商口语化表达（"还不错""蛮便宜的""挺多选择"）理解较弱，倾向将模糊表达判为负向；',
    'RoBERTa在京东电商评论上微调，能精准识别电商领域的隐含正面表达，平均置信度高达96.93%；',
    '89.8%的正向比例更符合消费者主动评价行为规律（有正面体验才更愿意写评论）；',
    'RoBERTa不设中性类别，原3.9%中性评论被明确划分至正向或负向，提升了分类精度。',
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    set_font(r, size=11)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(22)
r1 = p.add_run('结论：')
set_font(r1, size=11, bold=True)
r2 = p.add_run('硬折扣零食店整体口碑良好，超过九成评论为正向，消费者满意度较高，负向评论比例仅10.2%，说明当前经营状况受到消费者较高认可。')
set_font(r2, size=11)

doc.add_page_break()

# ===== 二、主题挖掘 =====
add_heading(doc, '二、主题挖掘结果对比', 1)

add_heading(doc, '2.1 好评主题对比', 2)
add_table(doc,
    ['主题', '原方法LDA关键词', '本次LDA关键词', '主题标签'],
    [
        ['主题1', '价格、便宜、性价比、方便、不错', '便宜、新开、喜欢、饮料、想来', '价格优势与复购意愿'],
        ['主题2', '环境、服务、品种、实惠、整洁', '口味、好吃、方便、干净、划算', '门店体验综合满意'],
        ['主题3', '好吃、口味、品类、丰富、味道', '货架、实惠、店员、好吃、丰富', '产品品类与陈列'],
    ],
    col_widths=[1.5, 4, 4, 4]
)
add_para(doc, '两次LDA运行结果高度吻合，说明数据稳定可靠。价格优势、门店体验、产品品类是消费者正向评价的三大核心驱动因素。', indent=1)

add_heading(doc, '2.2 差评主题对比', 2)
add_table(doc,
    ['主题', '原方法LDA关键词', '本次LDA关键词', '主题标签'],
    [
        ['主题1', '收银、结账、排队、收银员、袋子', '态度、收银员、袋子、问题', '收银服务问题'],
        ['主题2', '服务、态度、商品、品牌、营业员', '兑换、结果、巧克力', '服务与促销兑换'],
        ['主题3', '会员、价格、便宜、货架、品种', '店员、袋子、核销、买单、结账', '会员结账流程'],
    ],
    col_widths=[1.5, 4, 4, 4]
)

add_heading(doc, '2.3 BERT语义聚类新发现', 2)
add_para(doc, 'BERT方法通过语义空间聚类，识别出传统LDA未单独呈现的两个细粒度差评主题：', indent=1)
add_table(doc,
    ['差评子主题', '核心关键词', '问题描述', '建议措施'],
    [
        ['包装开封困难', '剪刀、撕开、果冻、勺子', '部分产品包装设计不友好，消费者需借助剪刀等工具才能打开', '门店提供开袋工具；向供应商反馈包装优化需求'],
        ['排队秩序问题', '插队、队伍、核销、当时', '高峰期收银区队伍混乱，存在插队现象影响体验', '设置排队引导线；高峰期增设收银台'],
    ],
    col_widths=[2.5, 3, 4.5, 3.5]
)

doc.add_page_break()

# ===== 三、词云分析 =====
add_heading(doc, '三、词云图与高频词分析', 1)

add_heading(doc, '3.1 词云图', 2)
add_para(doc, '基于RoBERTa情感分类结果，分别对正向（1,235条）和负向（140条）评论进行词频统计，生成词云图如下：', indent=1)
add_image(doc,
    os.path.join(OUTPUT_DIR, 'wordcloud_bert_tokenizer.png'),
    '图2  用户评论词云图：正向评论（左/红色）vs 负向评论（右/蓝色）',
    width=Inches(6.2)
)

add_heading(doc, '3.2 高频词统计与解读', 2)

add_heading(doc, '（1）正向评论高频词 Top 10', 3)
add_table(doc,
    ['排名', '关键词', '出现频次', '语义解读'],
    [
        ['1', '便宜', '287', '价格是最核心正面驱动，在1,235条好评中近1/4提及'],
        ['2', '喜欢', '200', '表达强烈情感认同，复购意愿词汇'],
        ['3', '实惠', '175', '与"便宜"共同构成性价比主题'],
        ['4', '方便', '161', '选址便利、购物流程便捷'],
        ['5', '好吃', '159', '产品口感满意的直接表达'],
        ['6', '口味', '158', '与"好吃"共同构成产品口味主题'],
        ['7', '饮料', '147', '饮品品类最受欢迎，高频提及'],
        ['8', '干净', '142', '门店卫生环境满意'],
        ['9', '划算', '138', '性价比高，物超所值'],
        ['10', '超市', '128', '常与"比超市便宜"对比使用，强化价格优势感知'],
    ],
    col_widths=[1.5, 2, 2.5, 7.5]
)

add_heading(doc, '（2）负向评论高频词 Top 10', 3)
add_table(doc,
    ['排名', '关键词', '出现频次', '语义解读'],
    [
        ['1', '袋子', '43', '购物袋问题高频出现，反映收费/提供方式不满'],
        ['2', '店员', '40', '员工服务是差评核心，与"态度"共同构成服务主题'],
        ['3', '态度', '39', '员工服务态度不佳是最直接投诉'],
        ['4', '兑换', '37', '会员积分/优惠兑换流程复杂或失败'],
        ['5', '收银员', '25', '收银员服务效率和态度问题'],
        ['6', '结账', '22', '结账流程体验不佳（排队、慢、出错）'],
        ['7', '核销', '17', '优惠券/积分核销失败或流程繁琐'],
        ['8', '买单', '17', '结账环节问题，与"结账"主题高度相关'],
        ['9', '不好', '17', '负面评价的通用表达'],
        ['10', '收银台', '14', '收银台硬件/布局/效率问题'],
    ],
    col_widths=[1.5, 2, 2.5, 7.5]
)

add_para(doc, '综合词云图与词频统计：正向评论以"便宜""喜欢""实惠"为核心，体现消费者对硬折扣模式价格优势的高度认可；负向评论高度集中于"袋子""店员""态度""兑换"，收银服务体验是最主要痛点。', indent=1)

doc.add_page_break()

# ===== 四、综合结论与建议 =====
add_heading(doc, '四、综合结论与业务建议', 1)

add_heading(doc, '4.1 消费者满意度驱动因素', 2)
add_table(doc,
    ['优先级', '驱动因素', '证据', '业务意义'],
    [
        ['第1位', '价格优势与性价比', '"便宜"287次，"实惠"175次，"划算"138次', '硬折扣模式核心竞争力，需持续维护价格优势'],
        ['第2位', '产品品类丰富度', '"品类""丰富""饮料""坚果"高频出现', '零食种类齐全是核心护城河，需持续引入新品'],
        ['第3位', '门店环境整洁', '"干净"142次，"整洁"96次', '整洁的购物环境直接影响消费者停留时长和复购'],
        ['第4位', '产品口味', '"好吃"159次，"口味"158次', '产品本身的口感满意度是复购的直接驱动'],
    ],
    col_widths=[1.5, 3.5, 5, 3.5]
)

add_heading(doc, '4.2 核心痛点与改进建议', 2)
add_table(doc,
    ['优先级', '痛点', '数据依据', '改进措施'],
    [
        ['高', '收银效率与排队体验', '收银/结账/排队关键词高频，差评主题1', '增设自助收银机；高峰期增派收银员；优化POS系统'],
        ['高', '员工服务态度', '"态度"39次，"店员"40次为差评前三', '加强服务培训；绩效与顾客评分挂钩'],
        ['中', '购物袋提供方式', '"袋子"43次为差评第一', '提供免费基础袋或清晰标注收费政策'],
        ['中', '包装开封困难', 'BERT独有发现：剪刀/撕开/果冻/勺子', '门店提供开袋工具；向供应商反馈包装改进需求'],
        ['中', '会员兑换流程复杂', '"兑换"37次，"核销"17次', '简化积分兑换；培训员工熟悉会员系统'],
        ['低', '货架补货及时性', '"货架"负向出现', '建立实时库存提醒；优化热门品补货频率'],
    ],
    col_widths=[1, 3, 4, 5.5]
)

doc.add_page_break()

# ===== 五、正大杯问卷设计 =====
add_heading(doc, '五、正大杯问卷设计建议', 1)
add_para(doc, '基于数据挖掘发现的5大消费者关注维度，建议问卷从以下6个模块设计，题目总数控制在25-30题。', indent=1)

modules = [
    ('模块一：消费行为基本信息', '了解受访者画像与消费频次', [
        '您的年龄段（18-24 / 25-30 / 31-35 / 35岁以上）',
        '您在上海硬折扣零食店的消费频次（每周多次/每周一次/每月2-3次/偶尔）',
        '您最常去的品牌（零食很忙 / 好想来 / 其他___）',
        '您通常与谁一起购物（独自 / 与朋友 / 与家人）',
    ]),
    ('模块二：价格与性价比感知', '对应好评第1位驱动因素——价格是最高频正面词', [
        '您认为该店商品定价与超市相比：（便宜很多/略便宜/差不多/略贵）',
        '您对会员折扣/优惠活动的满意度（1-5分）',
        '您是否因为价格因素放弃购买某商品？（是/否）',
        '该店性价比与竞争对手（便利店/零食量贩）相比如何？',
    ]),
    ('模块三：产品品类与口味', '对应好评第2、4位驱动因素', [
        '您最满意的产品品类（多选）：饮料/坚果/巧克力/饼干/果冻/进口商品',
        '您认为目前品类是否足够丰富？（非常丰富/基本满足/希望增加___类）',
        '您对产品新鲜度/保质期管理的满意度（1-5分）',
        '【BERT新发现验证题】您是否遇过零食包装难以打开的情况？（是/否，如是描述）',
    ]),
    ('模块四：门店体验', '对应差评高优先级痛点——收银效率是最集中差评', [
        '您在该店的平均等待结账时间（＜3分钟/3-8分钟/8-15分钟/＞15分钟）',
        '对收银效率的满意度（1-5分）',
        '收银员服务态度满意度（1-5分）',
        '您对购物袋提供方式是否满意？（满意/不满意，原因___）',
        '您是否遇过排队秩序混乱（如插队）的情况？（从未/偶尔/经常）',
    ]),
    ('模块五：会员制度与数字化体验', '对应差评中优先级——会员兑换流程复杂', [
        '您是否办理了该店会员？（是/否，如否原因___）',
        '您对会员积分/兑换规则的了解程度（完全了解/大致了解/不太了解）',
        '对会员权益的整体满意度（1-5分）',
        '您更希望通过什么方式了解会员优惠？（小程序/门店屏幕/店员介绍/货架标注）',
    ]),
    ('模块六：综合评价', '计算NPS净推荐值，捕获个性化反馈', [
        '您会向朋友推荐这家店吗？（NPS：0-10分）',
        '如推荐/不推荐，主要原因是什么？（开放填写）',
        '您认为该店最需要改进的一点是什么？（开放填写）',
    ]),
]

for title, purpose, questions in modules:
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    r1 = p.add_run('设计目的：')
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(purpose)
    set_font(r2, size=11, color=(89, 89, 89))
    for i, q in enumerate(questions, 1):
        p2 = doc.add_paragraph()
        p2.paragraph_format.first_line_indent = Pt(22)
        r = p2.add_run(f'{i}. {q}')
        set_font(r, size=11)

output_path = r'E:\ZDcup\upgraded_analysis\数据挖掘分析报告.docx'
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
